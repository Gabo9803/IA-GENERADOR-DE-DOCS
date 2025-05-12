from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, flash, Response
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from dotenv import load_dotenv
import openai
import os
import pdfplumber
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
from datetime import datetime
import re
import json
import sqlite3
import bcrypt
from collections import Counter
import statistics
import markdown
from docx import Document
import pytesseract
from PIL import Image
import hashlib

app = Flask(__name__)
app.secret_key = os.urandom(24)

# Cargar variables de entorno
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    raise ValueError("OPENAI_API_KEY no está configurada")

# Configurar cliente OpenAI
client = openai.OpenAI(api_key=openai_api_key)

# Configurar Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Modelo de usuario
class User(UserMixin):
    def __init__(self, id, email):
        self.id = id
        self.email = email

@login_manager.user_loader
def load_user(user_id):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('SELECT id, email FROM users WHERE id = ?', (user_id,))
        user = cursor.fetchone()
        if user:
            return User(user[0], user[1])
        return None

# Versión del esquema
CURRENT_SCHEMA_VERSION = 3

def init_db():
    """Inicializa la base de datos con las tablas necesarias."""
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS schema_version (
                id INTEGER PRIMARY KEY,
                version INTEGER NOT NULL
            )
        ''')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS style_profiles (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                name TEXT,
                fonts TEXT,
                font_size INTEGER,
                tone TEXT,
                margins TEXT,
                structure TEXT,
                text_color TEXT,
                background_color TEXT,
                alignment TEXT,
                line_spacing REAL,
                document_purpose TEXT,
                confidence_scores TEXT,
                analysis_keywords TEXT,
                embeddings TEXT,
                visual_elements TEXT,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')
        conn.commit()

def migrate_db():
    """Migra la base de datos a la versión actual."""
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('SELECT version FROM schema_version WHERE id = 1')
        result = cursor.fetchone()
        current_version = result[0] if result else 0

        if not result:
            cursor.execute('INSERT INTO schema_version (id, version) VALUES (1, ?)', (current_version,))
            conn.commit()

        if current_version < 1:
            current_version = 1

        if current_version < 2:
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN text_color TEXT DEFAULT "#000000"')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN background_color TEXT DEFAULT "#FFFFFF"')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN alignment TEXT DEFAULT "left"')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN line_spacing REAL DEFAULT 1.33')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN document_purpose TEXT DEFAULT "general"')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN confidence_scores TEXT DEFAULT "{}"')
            current_version = 2

        if current_version < 3:
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN name TEXT DEFAULT ""')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN fonts TEXT DEFAULT "{}"')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN analysis_keywords TEXT DEFAULT "[]"')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN embeddings TEXT DEFAULT "[]"')
            cursor.execute('ALTER TABLE style_profiles ADD COLUMN visual_elements TEXT DEFAULT "{}"')
            cursor.execute('UPDATE style_profiles SET fonts = ?', (json.dumps({'normal': 'Helvetica', 'bold': 'Helvetica-Bold', 'italic': 'Helvetica-Oblique'}),))
            current_version = 3

        cursor.execute('UPDATE schema_version SET version = ? WHERE id = 1', (CURRENT_SCHEMA_VERSION,))
        conn.commit()

init_db()
migrate_db()

# Cache con límite de tamaño
MAX_CACHE_SIZE = 100
response_cache = {}
session_messages = {}
html_previews = {}
analysis_cache = {}

def get_db_connection():
    """Abre una conexión a la base de datos con row_factory configurado."""
    conn = sqlite3.connect('profiles.db')
    conn.row_factory = sqlite3.Row
    return conn

def get_default_style_profile():
    """Devuelve un perfil de estilo por defecto."""
    return {
        'name': 'Default',
        'fonts': {'normal': 'Helvetica', 'bold': 'Helvetica-Bold', 'italic': 'Helvetica-Oblique'},
        'font_size': 12,
        'tone': 'neutral',
        'margins': {'top': 1, 'bottom': 1, 'left': 1, 'right': 1},
        'structure': ['paragraphs'],
        'text_color': '#000000',
        'background_color': '#FFFFFF',
        'alignment': 'left',
        'line_spacing': 1.33,
        'document_purpose': 'general',
        'confidence_scores': {'fonts': 1.0, 'font_size': 1.0, 'tone': 0.8, 'margins': 1.0, 'structure': 0.8, 'text_color': 1.0, 'background_color': 1.0, 'alignment': 0.9, 'line_spacing': 0.9, 'document_purpose': 0.7},
        'analysis_keywords': [],
        'embeddings': [],
        'visual_elements': {}
    }

def parse_markdown_to_reportlab(text, style_profile=None):
    """Convierte Markdown a elementos de ReportLab con soporte para estilos personalizados."""
    styles = getSampleStyleSheet()
    style_profile = style_profile or get_default_style_profile()
    fonts = style_profile['fonts']
    font_size = style_profile['font_size']
    alignment = {'left': 0, 'center': 1, 'right': 2, 'justified': 4}.get(style_profile['alignment'], 0)
    line_spacing = style_profile['line_spacing']
    text_color = colors.HexColor(style_profile['text_color'])
    background_color = colors.HexColor(style_profile['background_color'])

    # Validar fuentes
    available_fonts = pdfmetrics.getRegisteredFontNames()
    font_normal = fonts.get('normal', 'Helvetica')
    font_bold = fonts.get('bold', 'Helvetica-Bold')
    font_italic = fonts.get('italic', 'Helvetica-Oblique')

    if font_normal not in available_fonts:
        print(f"Fuente normal no válida '{font_normal}', usando 'Helvetica'")
        font_normal = 'Helvetica'
    if font_bold not in available_fonts:
        print(f"Fuente negrita no válida '{font_bold}', usando 'Helvetica-Bold'")
        font_bold = 'Helvetica-Bold'
    if font_italic not in available_fonts:
        print(f"Fuente cursiva no válida '{font_italic}', usando 'Helvetica-Oblique'")
        font_italic = 'Helvetica-Oblique'

    body_style = ParagraphStyle(
        name='Body',
        fontSize=font_size,
        leading=font_size * line_spacing,
        spaceAfter=8,
        fontName=font_normal,
        alignment=alignment,
        textColor=text_color,
        fbackColor=background_color
    )
    bold_style = ParagraphStyle(
        name='Bold',
        fontSize=font_size,
        leading=font_size * line_spacing,
        spaceAfter=8,
        fontName=font_bold,
        alignment=alignment,
        textColor=text_color,
        backColor=background_color
    )
    italic_style = ParagraphStyle(
        name='Italic',
        fontSize=font_size,
        leading=font_size * line_spacing,
        spaceAfter=8,
        fontName=font_italic,
        alignment=alignment,
        textColor=text_color,
        backColor=background_color
    )
    elements = []

    # Soporte para tablas Markdown
    table_pattern = re.compile(r'\|(.+?)\|\n\|[-|:\s]+\|\n((?:\|.+?\|(?:\n|$))+)')
    text_parts = table_pattern.split(text)

    for i, part in enumerate(text_parts):
        if i % 3 == 2:  # Contenido de la tabla
            rows = [row.strip().split('|')[1:-1] for row in part.strip().split('\n')]
            table_data = [[cell.strip() for cell in row] for row in rows]
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), font_normal),
                ('FONTSIZE', (0, 0), (-1, -1), font_size),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('LEADING', (0, 0), (-1, -1), font_size * line_spacing),
                ('TEXTCOLOR', (0, 0), (-1, -1), text_color),
                ('BACKCOLOR', (0, 0), (-1, -1), background_color),
            ]))
            elements.append(table)
            elements.append(Spacer(1, 0.2 * inch))
        elif i % 3 != 1:  # Texto normal
            lines = part.split('\n')
            for line in lines:
                line = line.replace('\n', '<br />')
                line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', line)
                if line.strip().startswith('- '):
                    line = f'• {line[2:]}'
                    elements.append(Paragraph(line, body_style))
                else:
                    style = bold_style if '<b>' in line else italic_style if '<i>' in line else body_style
                    elements.append(Paragraph(line, style))
                elements.append(Spacer(1, 0.1 * inch))

    return elements

def analyze_document(file):
    """Analiza un documento subido para extraer estilo, formato y tipografía con mayor precisión."""
    style_profile = {
        'name': f"Perfil - {datetime.now().strftime('%Y%m%d_%H%M%S')}",
        'fonts': {'normal': 'Helvetica', 'bold': 'Helvetica-Bold', 'italic': 'Helvetica-Oblique'},
        'font_size': 12,
        'tone': 'neutral',
        'margins': {'top': 1, 'bottom': 1, 'left': 1, 'right': 1},
        'structure': [],
        'text_color': '#000000',
        'background_color': '#FFFFFF',
        'alignment': 'left',
        'line_spacing': 1.33,
        'document_purpose': 'general',
        'confidence_scores': {
            'fonts': 0.8,
            'font_size': 0.8,
            'tone': 0.7,
            'margins': 0.8,
            'structure': 0.7,
            'text_color': 0.8,
            'background_color': 0.8,
            'alignment': 0.7,
            'line_spacing': 0.7,
            'document_purpose': 0.6
        },
        'analysis_keywords': [],
        'embeddings': [],
        'visual_elements': {}
    }

    content = ""
    try:
        # Calcular hash del archivo para caché
        file.seek(0)
        file_hash = hashlib.md5(file.read()).hexdigest()
        file.seek(0)
        if file_hash in analysis_cache:
            print(f"Usando caché para archivo con hash {file_hash}")
            return analysis_cache[file_hash]

        if file.filename.endswith('.pdf'):
            with pdfplumber.open(file) as pdf:
                font_names = []
                font_styles = []
                font_sizes = []
                margins = {'top': [], 'bottom': [], 'left': [], 'right': []}
                text_colors = []
                background_colors = []
                alignments = []
                line_spacings = []
                structures = set()
                heading_levels = []
                visual_elements = {'images': 0, 'tables': 0}

                max_pages = 5
                for page in pdf.pages[:min(len(pdf.pages), max_pages)]:
                    # Extraer texto
                    content += page.extract_text() or ""
                    print(f"Procesando página {page.page_number}, longitud del texto: {len(content)}")

                    # Analizar fuentes incrustadas
                    for font in page.fonts:
                        font_name = font.get('BaseFont', 'Helvetica').split('-')[0]
                        cleaned_name = re.sub(r'^[^a-zA-Z]+', '', font_name).lower()
                        if font.get('FontFile'):
                            try:
                                font_data = font['FontFile'].stream
                                pdfmetrics.registerFont(TTFont(cleaned_name, BytesIO(font_data)))
                                print(f"Registrada fuente personalizada: {cleaned_name}")
                            except Exception as e:
                                print(f"Error al registrar fuente {cleaned_name}: {e}")
                        font_names.append(cleaned_name)
                        font_styles.append(font.get('BaseFont', ''))

                    # Analizar caracteres
                    chars = page.chars
                    if chars:
                        raw_fonts = [char.get('fontname', 'Helvetica') for char in chars]
                        cleaned_fonts = [re.sub(r'^[^a-zA-Z]+', '', f).lower() for f in raw_fonts]
                        font_names.extend(cleaned_fonts)
                        font_sizes.extend([round(char.get('size', 12)) for char in chars])
                        text_colors.extend([char.get('non_stroking_color', (0, 0, 0)) for char in chars])

                    # Calcular márgenes
                    if page.chars:
                        margins['top'].append(page.bbox[3] - max(c['y1'] for c in page.chars))
                        margins['bottom'].append(min(c['y0'] for c in page.chars) - page.bbox[1])
                        margins['left'].append(min(c['x0'] for c in page.chars) - page.bbox[0])
                        margins['right'].append(page.bbox[2] - max(c['x1'] for c in page.chars))

                    # Detectar alineación y espaciado
                    if page.extract_text():
                        lines = page.extract_text().split('\n')
                        paragraph_spacings = []
                        for i in range(len(lines)-1):
                            try:
                                y0 = page.chars[i]['y0']
                                y1 = page.chars[i+1]['y0']
                                spacing = (y0 - y1) / page.chars[i].get('size', 12)
                                if spacing > 0 and not lines[i].strip().startswith(('-', '#', '|')):
                                    paragraph_spacings.append(spacing)
                            except IndexError:
                                continue
                        if paragraph_spacings:
                            line_spacings.append(statistics.mode(paragraph_spacings))

                        x_positions = [c['x0'] for c in page.chars]
                        if x_positions:
                            x_variance = statistics.variance(x_positions) if len(x_positions) > 1 else 0
                            if x_variance < 10:
                                alignments.append('justified' if max(x_positions) - min(x_positions) > page.width * 0.8 else 'left')
                            elif max(x_positions) > page.width * 0.9:
                                alignments.append('right')
                            elif abs(max(x_positions) + min(x_positions) - page.width) < page.width * 0.1:
                                alignments.append('center')

                    # Detectar estructuras avanzadas
                    avg_size = statistics.mean(font_sizes) if font_sizes else 12
                    for char in page.chars:
                        size = char.get('size', 12)
                        if size > avg_size * 1.2:
                            heading_levels.append({'size': size, 'text': char['text']})
                    if page.extract_tables():
                        structures.add('tables')
                        visual_elements['tables'] += len(page.extract_tables())
                    if page.extract_text():
                        text = page.extract_text()
                        if re.search(r'^\s*[\•\-\*]\s+', text, re.MULTILINE):
                            structures.add('lists')
                        if re.search(r'^\s*[A-Z].*\n[-=]+', text, re.MULTILINE):
                            structures.add('headings')
                        if re.search(r'\n\n+', text):
                            structures.add('paragraphs')
                        if any(line.strip().startswith('>') for line in text.split('\n')):
                            structures.add('blockquotes')
                        if any(re.match(r'^\d+\.\s+', line) for line in text.split('\n')[-5:]):
                            structures.add('footnotes')
                    if page.images:
                        visual_elements['images'] += len(page.images)

                # Normalizar valores
                if font_names:
                    font_counter = Counter(font_names)
                    raw_font = font_counter.most_common(1)[0][0]
                    print(f"Fuente detectada: {raw_font}")
                    font_mapping = {
                        'timesnewroman': 'Times-Roman',
                        'times new roman': 'Times-Roman',
                        'arial': 'Helvetica',
                        'arialm': 'Helvetica',
                        'helvetica': 'Helvetica',
                        'couriernew': 'Courier',
                        'courier new': 'Courier',
                        '': 'Helvetica'
                    }
                    style_profile['fonts']['normal'] = font_mapping.get(raw_font.lower(), 'Helvetica')
                    for style in font_styles:
                        if 'Bold' in style:
                            style_profile['fonts']['bold'] = font_mapping.get(raw_font.lower(), 'Helvetica') + '-Bold'
                        if 'Italic' in style or 'Oblique' in style:
                            style_profile['fonts']['italic'] = font_mapping.get(raw_font.lower(), 'Helvetica') + '-Oblique'
                    style_profile['confidence_scores']['fonts'] = font_counter.most_common(1)[0][1] / len(font_names)
                    print(f"Fuentes normalizadas: {style_profile['fonts']}")

                if not font_names:
                    style_profile['fonts'] = {'normal': 'Helvetica', 'bold': 'Helvetica-Bold', 'italic': 'Helvetica-Oblique'}
                    style_profile['confidence_scores']['fonts'] = 0.5
                    print("No se detectaron fuentes, usando fuentes por defecto")

                if font_sizes:
                    size_counter = Counter(font_sizes)
                    style_profile grove['font_size'] = size_counter.most_common(1)[0][0]
                    style_profile['confidence_scores']['font_size'] = size_counter.most_common(1)[0][1] / len(font_sizes)

                if margins['top']:
                    margin_values = {key: Counter([round(v/72, 2) for v in margins[key]]) for key in margins}
                    for key in margins:
                        style_profile['margins'][key] = margin_values[key].most_common(1)[0][0]
                        style_profile['confidence_scores']['margins'] = min(1.0, len(margins[key]) / len(pdf.pages))

                if text_colors:
                    color_counter = Counter([tuple(c) if isinstance(c, (list, tuple)) else c for c in text_colors])
                    dominant_color = color_counter.most_common(1)[0][0]
                    if isinstance(dominant_color, (list, tuple)):
                        style_profile['text_color'] = '#{:02x}{:02x}{:02x}'.format(
                            int(dominant_color[0] * 255), int(dominant_color[1] * 255), int(dominant_color[2] * 255)
                        )
                    style_profile['confidence_scores']['text_color'] = color_counter.most_common(1)[0][1] / len(text_colors)

                if background_colors:
                    bg_counter = Counter(background_colors)
                    dominant_bg = bg_counter.most_common(1)[0][0]
                    style_profile['background_color'] = '#{:02x}{:02x}{:02x}'.format(
                        int(dominant_bg[0] * 255), int(dominant_bg[1] * 255), int(dominant_bg[2] * 255)
                    )
                    style_profile['confidence_scores']['background_color'] = bg_counter.most_common(1)[0][1] / len(background_colors)

                if alignments:
                    alignment_counter = Counter(alignments)
                    style_profile['alignment'] = alignment_counter.most_common(1)[0][0]
                    style_profile['confidence_scores']['alignment'] = alignment_counter.most_common(1)[0][1] / len(alignments)

                if line_spacings:
                    style_profile['line_spacing'] = min(max(statistics.mean(line_spacings), 1.0), 2.0)
                    style_profile['confidence_scores']['line_spacing'] = len(line_spacings) / len(pdf.pages)

                if heading_levels:
                    unique_sizes = sorted(set(h['size'] for h in heading_levels), reverse=True)
                    style_profile['structure'].append({'headings': len(unique_sizes)})

                style_profile['structure'] = list(structures)
                style_profile['confidence_scores']['structure'] = 0.9 if structures else 0.7
                style_profile['visual_elements'] = visual_elements

        elif file.filename.endswith('.txt'):
            content = file.read().decode('utf-8', errors='ignore')
            md = markdown.Markdown()
            parsed = md.parse(content)
            structures = set()
            if any(tag.startswith('h') for tag in parsed):
                structures.add('headings')
            if 'ul' in parsed or 'ol' in parsed:
                structures.add('lists')
            if re.search(r'^\s*\|.*\|\s*$', content, re.MULTILINE):
                structures.add('tables')
            if re.search(r'^>\s+', content, re.MULTILINE):
                structures.add('blockquotes')
            if '```' in content:
                style_profile['fonts']['normal'] = 'Courier'
                style_profile['confidence_scores']['fonts'] = 0.7
            style_profile['structure'] = list(structures)
            style_profile['confidence_scores']['structure'] = 0.8 if structures else 0.6

        elif file.filename.endswith('.docx'):
            doc = Document(file)
            content = '\n'.join(p.text for p in doc.paragraphs)
            structures = set(['paragraphs'])
            for p in doc.paragraphs:
                if p.style.name.startswith('Heading'):
                    structures.add('headings')
                if p.style.name == 'List Bullet' or p.style.name == 'List Number':
                    structures.add('lists')
            if doc.tables:
                structures.add('tables')
            style_profile['structure'] = list(structures)
            style_profile['confidence_scores']['structure'] = 0.8 if structures else 0.6
            style_profile['visual_elements'] = {'tables': len(doc.tables)}

        elif file.filename.endswith(('.png', '.jpg', '.jpeg')):
            content = pytesseract.image_to_string(Image.open(file))
            structures = set(['paragraphs'])
            if re.search(r'^\s*[\•\-\*]\s+', content, re.MULTILINE):
                structures.add('lists')
            if re.search(r'^\s*[A-Z].*\n[-=]+', content, re.MULTILINE):
                structures.add('headings')
            style_profile['structure'] = list(structures)
            style_profile['confidence_scores']['structure'] = 0.7
            style_profile['visual_elements'] = {'images': 1}

        # Analizar tono y propósito
        if content.strip():
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "system",
                            "content": """
                            Analiza el texto y devuelve un JSON con:
                            - 'tone': tono (e.g., formal, informal, técnico, académico).
                            - 'document_purpose': propósito (e.g., informe, carta, manual, acta, propuesta).
                            - 'confidence': confianza (0.0 a 1.0).
                            - 'keywords': palabras clave del análisis.
                            """
                        },
                        {"role": "user", "content": content[:4000]}
                    ],
                    max_tokens=500
                )
                analysis = json.loads(response.choices[0].message.content)
                style_profile['tone'] = analysis.get('tone', 'neutral')
                style_profile['document_purpose'] = analysis.get('document_purpose', 'general')
                style_profile['confidence_scores']['tone'] = analysis.get('confidence', 0.7)
                style_profile['confidence_scores']['document_purpose'] = analysis.get('confidence', 0.7)
                style_profile['analysis_keywords'] = analysis.get('keywords', [])
                print(f"Análisis de tono completado: {style_profile['tone']}, propósito: {style_profile['document_purpose']}")
            except Exception as e:
                print(f"Error en análisis de tono: {e}")
                # Reglas de respaldo
                if 'acta de reunión' in content.lower():
                    style_profile['document_purpose'] = 'acta'
                    style_profile['confidence_scores']['document_purpose'] = 0.8
                elif re.search(r'^\s*Objetivo[s]?:', content, re.MULTILINE):
                    style_profile['document_purpose'] = 'propuesta'
                    style_profile['confidence_scores']['document_purpose'] = 0.8
                formal_keywords = ['estimado', 'cordialmente', 'respecto']
                technical_keywords = ['implementación', 'arquitectura', 'sistema']
                if any(kw in content.lower() for kw in formal_keywords):
                    style_profile['tone'] = 'formal'
                    style_profile['confidence_scores']['tone'] = 0.7
                elif any(kw in content.lower() for kw in technical_keywords):
                    style_profile['tone'] = 'technical'
                    style_profile['confidence_scores']['tone'] = 0.7
        else:
            style_profile['tone'] = 'neutral'
            style_profile['document_purpose'] = 'general'
            style_profile['confidence_scores']['tone'] = 0.5
            style_profile['confidence_scores']['document_purpose'] = 0.5

        # Generar embeddings
        if content.strip():
            try:
                embedding = client.embeddings.create(input=content[:4000], model="text-embedding-ada-002").data[0].embedding
                style_profile['embeddings'] = embedding[:100]  # Limitar tamaño
                print(f"Embeddings generados, longitud: {len(style_profile['embeddings'])}")
            except Exception as e:
                print(f"Error al generar embeddings: {e}")

        # Limitar caché
        if len(analysis_cache) >= MAX_CACHE_SIZE:
            analysis_cache.pop(next(iter(analysis_cache)))
        analysis_cache[file_hash] = (style_profile, content)
        print(f"Análisis completado para archivo {file.filename}, hash: {file_hash}")

        return style_profile, content

    except Exception as e:
        print(f"Error al analizar documento: {e}")
        return get_default_style_profile(), ""

@app.route('/')
@login_required
def index():
    """Renderiza la página principal."""
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Maneja el inicio de sesión del usuario."""
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT id, email, password FROM users WHERE email = ?', (email,))
            user = cursor.fetchone()

            if user and bcrypt.checkpw(password.encode('utf-8'), user['password'].encode('utf-8')):
                login_user(User(user['id'], user['email']))
                return redirect(url_for('index'))
            else:
                flash('Correo o contraseña incorrectos', 'error')

    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    """Maneja el registro de nuevos usuarios."""
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('INSERT INTO users (email, password) VALUES (?, ?)', (email, hashed_password))
                conn.commit()
                flash('Registro exitoso. Por favor, inicia sesión.', 'success')
                return redirect(url_for('login'))
        except sqlite3.IntegrityError:
            flash('El correo ya está registrado', 'error')

    return render_template('register.html')

@app.route('/logout')
@login_required
def logout():
    """Cierra la sesión del usuario."""
    logout_user()
    return redirect(url_for('login'))

@app.route('/generate', methods=['POST'])
@login_required
def generate_document():
    """Genera un documento basado en el prompt y configuración proporcionada."""
    try:
        data = request.json
        prompt = data.get('prompt', '')
        session_id = data.get('session_id', 'default')
        doc_type = data.get('doc_type', 'general')
        tone = data.get('tone', 'neutral')
        length = data.get('length', 'medium')
        language = data.get('language', 'es')
        style_profile_id = data.get('style_profile_id', None)
        message_history = data.get('message_history', [])

        if not prompt:
            return jsonify({'error': 'El prompt está vacío'}), 400

        cache_key = f"{current_user.id}:{session_id}:{prompt}:{doc_type}:{tone}:{length}:{language}:{style_profile_id}"
        if cache_key in response_cache:
            print(f"Usando caché para clave {cache_key}")
            return jsonify({
                'document': response_cache[cache_key],
                'content_type': 'html' if doc_type == 'html' or 'html' in prompt.lower() else 'pdf'
            })

        if session_id not in session_messages:
            session_messages[session_id] = []

        style_profile = get_style_profile(style_profile_id)
        print(f"style_profile en /generate: {style_profile}")

        effective_tone = style_profile['tone'] if style_profile.get('tone') else tone
        structure = style_profile['structure'] if style_profile.get('structure') else ['paragraphs']
        structure_instruction = f"Usa una estructura que incluya: {', '.join([s if isinstance(s, str) else f'headings ({s.get('headings')} niveles)' for s in structure])}."

        is_explanatory = re.search(r'^(¿Quién es|¿Qué es|explicar|detallar)\b', prompt, re.IGNORECASE) is not None
        is_html = doc_type == 'html' or re.search(r'\b(html|página web|sitio web)\b', prompt, re.IGNORECASE) is not None

        if is_explanatory:
            system_prompt = f"""
            Eres GarBotGPT, un asistente que genera documentos profesionales en formato Markdown con una estructura clara para explicaciones.
            - Tipo de documento: {doc_type} (e.g., explicación, biografía, informe).
            - Tono: {effective_tone} (formal, informal, técnico).
            - Longitud: {length} (corto: ~100 palabras, medio: ~300 palabras, largo: ~600 palabras).
            - Idioma: {language} (e.g., es para español, en para inglés, fr para francés).
            - {structure_instruction}
            - Usa la siguiente estructura en Markdown:
              ```markdown
              # [Tema o Nombre]
              
              ## Introducción
              [Párrafo breve presentando el tema o persona.]
              
              ## Detalles Principales
              - [Clave 1]: [Valor o descripción]
              - [Clave 2]: [Valor o descripción]
              - ...
              
              ## Contexto Adicional
              [Párrafos detallando información relevante, como antecedentes, logros, o impacto.]
              
              ## Conclusión
              [Resumen de la relevancia o importancia del tema.]
              ```
            - Usa encabezados (#, ##), listas (-), negritas (**), y tablas (|...|) cuando sea apropiado.
            - Considera el contexto de los mensajes anteriores para mantener coherencia en la conversación.
            """
        elif is_html:
            system_prompt = f"""
            Eres GarBotGPT, un asistente que genera páginas web en formato HTML con CSS y JavaScript si es necesario.
            - Tipo de documento: página web (HTML).
            - Tono: {effective_tone} (formal, informal, técnico).
            - Longitud: {length} (corto: ~100 líneas, medio: ~300 líneas, largo: ~600 líneas).
            - Idioma: {language} (e.g., es para español, en para inglés, fr para francés).
            - Genera código HTML completo con:
              - Estructura semántica (<header>, <main>, <footer>, etc.).
              - Estilos CSS internos (en <style>) adaptados al tono y propósito.
              - JavaScript opcional (en <script>) si el prompt lo requiere.
              - Usa un diseño responsivo y moderno.
            - Considera el contexto de los mensajes anteriores para mantener coherencia.
            """
        else:
            system_prompt = f"""
            Eres GarBotGPT, un asistente que genera documentos profesionales en formato Markdown.
            - Tipo de documento: {doc_type} (e.g., carta formal, informe, correo, contrato, currículum).
            - Tono: {effective_tone} (formal, informal, técnico).
            - Longitud: {length} (corto: ~100 palabras, medio: ~300 palabras, largo: ~600 palabras).
            - Idioma: {language} (e.g., es para español, en para inglés, fr para francés).
            - {structure_instruction}
            - Usa encabezados (#, ##), listas (-), negritas (**), y tablas (|...|) cuando sea apropiado.
            - Considera el contexto de los mensajes anteriores para mantener coherencia en la conversación.
            """

        messages = [{"role": "system", "content": system_prompt}]
        messages.extend(message_history[-10:])
        messages.append({"role": "user", "content": prompt})

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=1000,
            temperature=0.7
        )

        document = response.choices[0].message.content

        # Limitar caché
        if len(response_cache) >= MAX_CACHE_SIZE:
            response_cache.pop(next(iter(response_cache)))
        response_cache[cache_key] = document

        session_messages[session_id].append({"role": "user", "content": prompt})
        session_messages[session_id].append({"role": "assistant", "content": document})

        return jsonify({
            'document': document,
            'content_type': 'html' if is_html else 'pdf'
        })
    except openai.AuthenticationError:
        print("Error de autenticación con OpenAI")
        return jsonify({'error': 'Error de autenticación con OpenAI. Verifica la clave API.'}), 401
    except openai.RateLimitError:
        print("Límite de solicitudes alcanzado en OpenAI")
        return jsonify({'error': 'Límite de solicitudes alcanzado. Intenta de nuevo más tarde.'}), 429
    except Exception as e:
        print(f"Error en /generate: {str(e)}")
        return jsonify({'error': f'Error al generar el documento: {str(e)}'}), 500

def get_style_profile(profile_id):
    """Obtiene un perfil de estilo o devuelve el estilo por defecto."""
    if not profile_id:
        return get_default_style_profile()
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM style_profiles WHERE id = ? AND user_id = ?', (profile_id, current_user.id))
        profile = cursor.fetchone()
        if profile:
            return {
                'name': profile['name'],
                'fonts': json.loads(profile['fonts']),
                'font_size': profile['font_size'],
                'tone': profile['tone'],
                'margins': json.loads(profile['margins']),
                'structure': json.loads(profile['structure']),
                'text_color': profile['text_color'],
                'background_color': profile['background_color'],
                'alignment': profile['alignment'],
                'line_spacing': profile['line_spacing'],
                'document_purpose': profile['document_purpose'],
                'confidence_scores': json.loads(profile['confidence_scores']),
                'analysis_keywords': json.loads(profile['analysis_keywords']),
                'embeddings': json.loads(profile['embeddings']),
                'visual_elements': json.loads(profile['visual_elements'])
            }
    return get_default_style_profile()

@app.route('/generate_preview', methods=['POST'])
@login_required
def generate_preview():
    """Genera una vista previa del documento en PDF o HTML."""
    try:
        data = request.json
        content = data.get('content', '')
        content_type = data.get('content_type', 'pdf')
        style_profile_id = data.get('style_profile_id', None)
        if not content:
            return jsonify({'error': 'El contenido está vacío'}), 400

        if content_type == 'html':
            preview_id = f"{current_user.id}:{datetime.now().timestamp()}"
            if len(html_previews) >= MAX_CACHE_SIZE:
                html_previews.pop(next(iter(html_previews)))
            html_previews[preview_id] = content
            return jsonify({'preview_id': preview_id, 'content_type': 'html'})

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=inch,
            bottomMargin=0.75*inch
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(name='Title', fontSize=18, leading=22, spaceAfter=12, fontName='Helvetica-Bold')
        subtitle_style = ParagraphStyle(name='Subtitle', fontSize=10, leading=14, spaceAfter=10, fontName='Helvetica-Oblique')

        story = []
        story.append(Paragraph("Documento Generado por GarBotGPT", title_style))
        story.append(Paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", subtitle_style))
        story.append(Spacer(1, 0.3 * inch))

        style_profile = get_style_profile(style_profile_id)
        print(f"Aplicando style_profile en generate_preview: {style_profile}")
        story.extend(parse_markdown_to_reportlab(content, style_profile))

        doc.build(story)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='documento_generado.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        print(f"Error en /generate_preview: {str(e)}")
        return jsonify({'error': f'Error al generar la vista previa: {str(e)}'}), 500

@app.route('/preview_html/<preview_id>', methods=['GET'])
@login_required
def preview_html(preview_id):
    """Devuelve el contenido HTML almacenado para la vista previa."""
    if preview_id in html_previews and preview_id.startswith(f"{current_user.id}:"):
        return Response(html_previews[preview_id], mimetype='text/html')
    return jsonify({'error': 'Vista previa no encontrada'}), 404

@app.route('/analyze_document', methods=['POST'])
@login_required
def analyze_document_endpoint():
    """Analiza un documento subido y guarda el perfil de estilo."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No se proporcionó ningún archivo'}), 400

        file = request.files['file']
        if not (file.filename.endswith(('.pdf', '.txt', '.docx', '.png', '.jpg', '.jpeg'))):
            return jsonify({'error': 'Formato no soportado. Usa PDF, TXT, DOCX, PNG o JPG.'}), 400

        style_profile, content = analyze_document(file)

        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO style_profiles (
                    user_id, name, fonts, font_size, tone, margins, structure,
                    text_color, background_color, alignment, line_spacing,
                    document_purpose, confidence_scores, analysis_keywords,
                    embeddings, visual_elements
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                current_user.id,
                style_profile['name'],
                json.dumps(style_profile['fonts']),
                style_profile['font_size'],
                style_profile['tone'],
                json.dumps(style_profile['margins']),
                json.dumps(style_profile['structure']),
                style_profile['text_color'],
                style_profile['background_color'],
                style_profile['alignment'],
                style_profile['line_spacing'],
                style_profile['document_purpose'],
                json.dumps(style_profile['confidence_scores']),
                json.dumps(style_profile['analysis_keywords']),
                json.dumps(style_profile['embeddings']),
                json.dumps(style_profile['visual_elements'])
            ))
            profile_id = cursor.lastrowid
            conn.commit()

        summary = {
            'content_preview': content[:500],
            'word_count': len(content.split()),
            'page_count': getattr(analyze_document, 'page_count', 1),
            'detected_structures': style_profile['structure'],
            'tone_analysis': {
                'tone': style_profile['tone'],
                'confidence': style_profile['confidence_scores']['tone'],
                'keywords': style_profile['analysis_keywords']
            },
            'visual_elements': style_profile['visual_elements']
        }

        return jsonify({
            'style_profile_id': str(profile_id),
            'style_profile': style_profile,
            'summary': summary
        })
    except Exception as e:
        print(f"Error en /analyze_document: {str(e)}")
        return jsonify({'error': f'Error al analizar el documento: {str(e)}'}), 500

@app.route('/style_profiles', methods=['GET'])
@login_required
def get_style_profiles():
    """Obtiene todos los perfiles de estilo del usuario."""
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM style_profiles WHERE user_id = ?', (current_user.id,))
        profiles = cursor.fetchall()

        result = {}
        for profile in profiles:
            result[str(profile['id'])] = {
                'name': profile['name'],
                'fonts': json.loads(profile['fonts']),
                'font_size': profile['font_size'],
                'tone': profile['tone'],
                'margins': json.loads(profile['margins']),
                'structure': json.loads(profile['structure']),
                'text_color': profile['text_color'],
                'background_color': profile['background_color'],
                'alignment': profile['alignment'],
                'line_spacing': profile['line_spacing'],
                'document_purpose': profile['document_purpose'],
                'confidence_scores': json.loads(profile['confidence_scores']),
                'analysis_keywords': json.loads(profile['analysis_keywords']),
                'embeddings': json.loads(profile['embeddings']),
                'visual_elements': json.loads(profile['visual_elements'])
            }

        return jsonify(result)

@app.route('/style_profiles/<profile_id>', methods=['DELETE'])
@login_required
def delete_style_profile(profile_id):
    """Elimina un perfil de estilo específico."""
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('DELETE FROM style_profiles WHERE id = ? AND user_id = ?', (profile_id, current_user.id))
        conn.commit()

        if cursor.rowcount > 0:
            return jsonify({'message': 'Perfil de estilo eliminado'})
        return jsonify({'error': 'Perfil no encontrado'}), 404

@app.route('/purge_style_profiles', methods=['DELETE'])
@login_required
def purge_style_profiles():
    """Elimina todos los perfiles de estilo del usuario."""
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM style_profiles WHERE user_id = ?', (current_user.id,))
            conn.commit()
        return jsonify({'message': 'Todos los perfiles de estilo han sido eliminados'})
    except Exception as e:
        print(f"Error en /purge_style_profiles: {str(e)}")
        return jsonify({'error': f'Error al purgar perfiles: {str(e)}'}), 500

if __name__ == '__main__':
    # Configuración para pruebas locales
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=True)